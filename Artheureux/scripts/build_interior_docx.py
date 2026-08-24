#!/usr/bin/env python3
"""Build the editable KDP interior DOCX for one Artheureux edition."""
from __future__ import annotations

import argparse
import json
import re
import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
CONFIG = {
    "adult": {
        "data": ROOT / "Artheureux/data/adult.json",
        "frontmatter": ROOT / "Artheureux/manuscript/adult-frontmatter.md",
        "label": "Édition adulte",
        "subtitle": "Cinquante règles pratiques inspirées de la pensée de Schopenhauer",
        "trim": (6.0, 9.0),
        "body_size": 10.2,
        "image_width": 4.0,
        "accent": "235B4E",
        "accent_soft": "E5EFEA",
        "key_fill": "173B33",
    },
    "youth": {
        "data": ROOT / "Artheureux/data/youth.json",
        "frontmatter": ROOT / "Artheureux/manuscript/youth-frontmatter.md",
        "label": "Édition 10–15 ans",
        "subtitle": "Cinquante idées pour réfléchir à ce que tu vis",
        "trim": (7.0, 10.0),
        "body_size": 11.0,
        "image_width": 4.1,
        "accent": "176B6B",
        "accent_soft": "E5F2EF",
        "key_fill": "18545A",
    },
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def add_mirror_margins(document: Document) -> None:
    settings = document.settings._element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))


def set_run_font(run, name: str, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document, config: dict) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(config["trim"][0])
    section.page_height = Inches(config["trim"][1])
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.58)
    section.gutter = Inches(0.08)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True
    add_mirror_margins(document)

    normal = document.styles["Normal"]
    normal.font.name = "Noto Serif"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif")
    normal.font.size = Pt(config["body_size"])
    normal.font.color.rgb = RGBColor.from_string("26302A")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (
        ("Title", 30, "17231D"),
        ("Heading 1", 22, "17231D"),
        ("Heading 2", 15, config["accent"]),
        ("Heading 3", 11, config["accent"]),
    ):
        style = document.styles[name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Georgia")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("L’ART D’ÊTRE HEUREUX  ·  DAVID DEVESA")
    set_run_font(header_run, "Noto Sans", 7.2, True, "68736C")

    footer_p = section.footer.paragraphs[0]
    add_page_field(footer_p)
    for run in footer_p.runs:
        set_run_font(run, "Noto Sans", 7.5, False, "68736C")


def add_title_page(document: Document, config: dict) -> None:
    document.add_paragraph("L’ART D’ÊTRE HEUREUX", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    set_run_font(subtitle.add_run(config["subtitle"]), "Georgia", 15, False, config["accent"], True)

    edition = document.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.paragraph_format.space_before = Pt(120)
    set_run_font(edition.add_run(config["label"]), "Noto Sans", 10, True, "68736C")

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(18)
    set_run_font(author.add_run("David DEVESA"), "Georgia", 17, False, "17231D")
    document.add_page_break()


def add_copyright_page(document: Document, config: dict) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(220)
    lines = [
        "© 2026 David DEVESA. Tous droits réservés.",
        f"{config['label']} — nouvelle édition.",
        "Adaptation pratique originale à partir de la pensée et de sources publiques de Schopenhauer.",
        "Les phrases clés et les exemples de ce livre ne sont pas des citations de Schopenhauer.",
        "ISBN : à attribuer pour cette nouvelle édition.",
    ]
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(line), "Noto Sans", 8.3, False, "68736C")
    document.add_page_break()


def strip_markdown(value: str) -> str:
    return re.sub(r"\*\*(.*?)\*\*", r"\1", value).replace("*", "").strip()


def add_reader_paths_table(document: Document, paths: dict, page_map: dict[str, int]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Thème", "Situation", "Règles et pages")
    for cell, label in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "DCEBE4")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(label), "Noto Sans", 8, True, "173B33")
    set_repeat_table_header(table.rows[0])

    for category in paths["categories"]:
        start_row = len(table.rows)
        for entry in category["entries"]:
            cells = table.add_row().cells
            # Repeat the category on every row. Large categories can span a
            # printed page; a vertically merged cell would leave continuation
            # pages with an unexplained blank first column.
            values = [category["label"], entry["label"]]
            pairs = []
            for rule_id in entry["rule_ids"]:
                number = int(rule_id.split("-")[-1])
                page = page_map.get(str(number)) or page_map.get(rule_id)
                pairs.append(f"{number} (p. {page})" if page else str(number))
            values.append(", ".join(pairs))
            for cell, value in zip(cells, values):
                set_cell_margins(cell, top=90, bottom=90)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                set_run_font(paragraph.add_run(value), "Noto Sans", 7.2, False, "26302A")
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cells[0], "EEF3ED")
    document.add_paragraph()


def add_frontmatter(document: Document, path: Path, paths: dict, page_map: dict[str, int]) -> None:
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "{{READER_PATHS_TABLE_FINAL_PAGINATION}}":
            add_reader_paths_table(document, paths, page_map)
            continue
        if page_map and line.startswith("*Le tableau et les numéros de page"):
            continue
        if line.startswith("# "):
            document.add_heading(strip_markdown(line[2:]), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(strip_markdown(line[3:]), level=2)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.24)
            paragraph.paragraph_format.first_line_indent = Inches(-0.24)
            paragraph.add_run(f"{numbered.group(1)}.  {strip_markdown(numbered.group(2))}")
            continue
        paragraph = document.add_paragraph()
        if line.startswith("*") and line.endswith("*"):
            set_run_font(paragraph.add_run(strip_markdown(line)), "Noto Serif", None, False, "68736C", True)
        else:
            paragraph.add_run(strip_markdown(line))
    document.add_page_break()


def add_label(document: Document, text: str, config: dict) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(text.upper()), "Noto Sans", 8.2, True, config["accent"])


def add_numbered_items(document: Document, items: list[str]) -> None:
    # Write visible numbers ourselves instead of relying on Word's shared list
    # counters. LibreOffice otherwise continues numbering across unrelated
    # sections and even across rules (for example 43, 44, 45).
    for number, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.24)
        paragraph.paragraph_format.first_line_indent = Inches(-0.24)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(f"{number}.  {item}")


def add_key_phrase(document: Document, phrase: str, config: dict) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, config["key_fill"])
    set_cell_margins(cell, top=220, start=260, bottom=220, end=260)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(phrase), "Georgia", 14.5, False, "FFFFFF", True)
    document.add_paragraph()


def rule_image_path(rule: dict, profile: str, image_cache: Path) -> Path:
    image = rule["image"]
    if profile == "print":
        return ROOT / "Artheureux" / image["print_path"]

    source = ROOT / "Artheureux" / image["web_path"]
    if not source.exists():
        return source
    target = image_cache / f"rule-{rule['order']:02d}.jpg"
    with Image.open(source) as opened:
        opened.convert("RGB").save(target, "JPEG", quality=80, optimize=True, progressive=True)
    return target


def add_rule(document: Document, rule: dict, edition: str, config: dict,
             image_profile: str, image_cache: Path) -> None:
    document.add_page_break()
    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.keep_with_next = True
    set_run_font(
        eyebrow.add_run(f"RÈGLE {rule['order']:02d}  ·  {config['label'].upper()}"),
        "Noto Sans", 8, True, config["accent"],
    )
    title = document.add_heading(rule["title"], level=1)
    title.paragraph_format.space_before = Pt(0)

    image_path = rule_image_path(rule, image_profile, image_cache)
    if image_path.exists():
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_after = Pt(10)
        picture.add_run().add_picture(str(image_path), width=Inches(config["image_width"]))

    add_label(document, "Résumé de la règle", config)
    document.add_paragraph(rule["summary"])

    add_label(document, "Exemple pratique", config)
    example = document.add_paragraph(rule["example"])
    example.paragraph_format.left_indent = Inches(0.12)
    example.paragraph_format.right_indent = Inches(0.12)

    add_label(document, "Conseils d’application", config)
    add_numbered_items(document, rule["advice"])

    add_key_phrase(document, rule["key_phrase"], config)

    add_label(document, "Questions à se poser", config)
    add_numbered_items(document, rule["questions"])

    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(12)
    source.paragraph_format.keep_together = True
    set_run_font(source.add_run("Point de départ philosophique — "), "Noto Sans", 7.2, True, "68736C")
    set_run_font(
        source.add_run(f"{rule['source']['work']} · {rule['source']['section']}"),
        "Noto Sans", 7.2, False, "68736C",
    )
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(0)
    set_run_font(
        note.add_run("Adaptation originale ; la phrase clé n’est pas une citation de Schopenhauer."),
        "Noto Sans", 7.0, False, "68736C", True,
    )


def add_contents(document: Document, rules: list[dict], page_map: dict[str, int], config: dict) -> None:
    document.add_heading("Les cinquante règles", level=1)
    for rule in rules:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(config["trim"][0] - 1.65))
        set_run_font(paragraph.add_run(f"{rule['order']:02d}. {rule['title']}"), "Noto Serif", 8.4, False, "26302A")
        page = page_map.get(str(rule["order"])) or page_map.get(rule["id"])
        if page:
            set_run_font(paragraph.add_run(f"\t{page}"), "Noto Sans", 8, True, config["accent"])
    document.add_page_break()


def load_page_map(path: str | None) -> dict[str, int]:
    if not path:
        return {}
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    return {str(key): int(value) for key, value in payload.items()}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--edition", choices=sorted(CONFIG), required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--page-map")
    parser.add_argument("--image-profile", choices=("print", "web"), default="print")
    args = parser.parse_args()

    config = CONFIG[args.edition]
    content = json.loads(config["data"].read_text(encoding="utf-8"))
    paths = json.loads((ROOT / "Artheureux/data/reader-paths.json").read_text(encoding="utf-8"))
    page_map = load_page_map(args.page_map)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix=f"artheureux-{args.edition}-images-") as temporary:
        document = Document()
        configure_document(document, config)
        add_title_page(document, config)
        add_copyright_page(document, config)
        add_frontmatter(document, config["frontmatter"], paths, page_map)
        add_contents(document, content["rules"], page_map, config)
        image_cache = Path(temporary)
        for rule in content["rules"]:
            add_rule(document, rule, args.edition, config, args.image_profile, image_cache)
        document.save(output)
    print(
        f"WROTE {output} rules={len(content['rules'])} edition={args.edition} "
        f"images={args.image_profile}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
